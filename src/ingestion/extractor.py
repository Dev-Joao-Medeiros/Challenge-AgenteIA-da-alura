from pathlib import Path
from pypdf import PdfReader
from pdf2image import convert_from_path
import pytesseract
from docx import Document 
import pandas as pd 
from tabulate import tabulate
import fitz  # PyMuPDF

# TESTE PARA PEGAR OS PARAGRAFOS COM FITZ
def extrair_pdf_test(caminho_arquivo: Path) -> tuple[str, dict]:
    documento = fitz.open(caminho_arquivo)

    textos_por_pagina = []
    ocr_foi_usado = False

    for pagina in documento:
        texto_pagina = pagina.get_text("text")

        # Mesma lógica de antes: se o texto extraído for muito curto/vazio,
        # a página provavelmente é uma imagem (PDF escaneado) -- aciona
        # o OCR só para essa página.
        if len(texto_pagina.strip()) < 20:
            imagem = pagina.get_pixmap().pil_image()
            texto_pagina = pytesseract.image_to_string(imagem, lang="por")
            ocr_foi_usado = True

        textos_por_pagina.append(texto_pagina)

    metadados = {
        "paginas": documento.page_count,
        "ocr": ocr_foi_usado,
    }

    documento.close()

    texto = "\n\n".join(textos_por_pagina)

    return texto, metadados

# ORIGINAL PYPDF PARA EXTRAIR OS TEXTO
def extrair_pdf(caminho_arquivo: Path) -> tuple[str, dict]:
    reader = PdfReader(caminho_arquivo)

    textos_por_pagina = []
    ocr_foi_usado = False

    for indice, pagina in enumerate(reader.pages):
        texto_pagina = pagina.extract_text() or ""

        # Se o texto extraído for muito curto/vazio, a página provavelmente
        # é uma imagem (PDF escaneado) -- aciona o OCR só para essa página.
        if len(texto_pagina.strip()) < 20:
            imagem = converter_pagina_em_imagem(caminho_arquivo, indice)
            texto_pagina = pytesseract.image_to_string(imagem, lang="por")
            ocr_foi_usado = True

        textos_por_pagina.append(texto_pagina)

    texto = "\n".join(textos_por_pagina)

    metadados = {
        "paginas": len(reader.pages),
        "ocr": ocr_foi_usado,
    }

    return texto, metadados


def converter_pagina_em_imagem(caminho_arquivo: Path, indice_pagina: int):
    imagens = convert_from_path(
        caminho_arquivo,
        first_page=indice_pagina + 1,
        last_page=indice_pagina + 1,
    )
    return imagens[0]

def extrair_docx(caminho_arquivo: Path) -> tuple[str, dict]:

    document = Document(caminho_arquivo)
    partes_texto = []
    contador_titulos = 0

    for paragrafo in document.paragraphs:
        texto_paragrafo = paragrafo.text.strip()

        if not texto_paragrafo:
            continue
        if paragrafo.style.name.startswith("Heading"):
            contador_titulos +=1 
            partes_texto.append(f"\n## {texto_paragrafo}\n")
        else:
            partes_texto.append(texto_paragrafo)

    texto = "\n".join(partes_texto)

    metadados = {
        "paragrafo" : len(document.paragraphs),
        "titulos_encontrados" : contador_titulos, 
    }
    
    return texto, metadados

def extrair_xlsx(caminho_do_arquivo: Path) -> tuple[str, dict]:
    abas = pd.read_excel(caminho_do_arquivo, sheet_name=None)

    texto_completo = []
    resumo_abas  = []

    for nome_da_aba, df_aba in abas.items():
        df_limpo = df_aba.fillna("")
        
        texto_markdown = df_limpo.to_markdown(index=False)
        
        bloco_texto = f"--- TEXTO DA ABA: {nome_da_aba} ---\n{texto_markdown}\n\n"
        texto_completo.append(bloco_texto)

        resumo_abas .append({
            "nome": nome_da_aba,
            "linhas": len(df_limpo),
            "colunas": len(df_limpo.columns),
        })
        
    string_final = "".join(texto_completo)

    metadados = {
        "total_abas": len(abas),
        "abas": resumo_abas,
    }

    return string_final, metadados

def extrair_csv(caminho_do_arquivo: Path) -> tuple[str, dict]:
    df = pd.read_csv(caminho_do_arquivo)
    
    df_limpo = df.fillna("")
    
    texto_markdown = df_limpo.to_markdown(index=False)
    
    nome_do_arquivo = caminho_do_arquivo.name
    string_final = f"--- TEXTO DO ARQUIVO: {nome_do_arquivo} ---\n{texto_markdown}\n\n"

    metadados = {
        "arquivo": nome_do_arquivo,
        "linhas": len(df_limpo),
        "colunas": len(df_limpo.columns),
    }

    return string_final, metadados

EXTRATORES_POR_EXTENSAO = {
    ".pdf": extrair_pdf_test, # VERIFICAR SE VAI USAR O extrair_pdf ou o extrair_pdf_test
    ".docx": extrair_docx,
    ".xlsx": extrair_xlsx,
    ".csv": extrair_csv,
}

def extrair(caminho_do_arquivo: Path) -> tuple[str,dict]:
    extensao = caminho_do_arquivo.suffix.lower()

    extrator = EXTRATORES_POR_EXTENSAO.get(extensao)

    if not extrator:
        extensoes_suportadas = " , ".join(EXTRATORES_POR_EXTENSAO.keys())
        raise ValueError(
            f"Extensão '{extensao}' não suportada. "
            f"Formatos aceitos: {extensoes_suportadas}"
        )

    return extrator(caminho_do_arquivo)